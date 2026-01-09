"""Extract text content from uploaded files (PDF, DOC, DOCX)."""

import io
from pathlib import Path
from typing import BinaryIO

from pypdf import PdfReader
from docx import Document


def extract_text_from_pdf(file: BinaryIO) -> str:
    """Extract text content from a PDF file."""
    try:
        reader = PdfReader(file)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(file: BinaryIO) -> str:
    """Extract text content from a DOCX file."""
    try:
        doc = Document(file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_file(filename: str, file_content: bytes) -> str:
    """Extract text from a file based on its extension."""
    ext = Path(filename).suffix.lower()
    file_io = io.BytesIO(file_content)

    if ext == ".pdf":
        return extract_text_from_pdf(file_io)
    elif ext in (".docx", ".doc"):
        # Note: .doc (old format) is not supported by python-docx
        # Only .docx works
        if ext == ".doc":
            return "[Error: Old .doc format not supported. Please convert to .docx]"
        return extract_text_from_docx(file_io)
    else:
        return f"[Unsupported file type: {ext}]"


def extract_text_from_files(files: list[tuple[str, bytes]]) -> str:
    """Extract text from multiple files and combine them.

    Args:
        files: List of (filename, content) tuples

    Returns:
        Combined text from all files with source labels
    """
    if not files:
        return ""

    extracted_parts = []
    for filename, content in files:
        text = extract_text_from_file(filename, content)
        if text and not text.startswith("[Error"):
            extracted_parts.append(f"--- Content from {filename} ---\n{text}")

    return "\n\n".join(extracted_parts)
